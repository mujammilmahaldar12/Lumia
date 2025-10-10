#!/usr/bin/env python3
"""
Lumia Black Book Generator - Convert TXT chapters to DOCX
Creates a professionally formatted Word document from chapter text files
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
import os
import glob

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    
    # Create custom styles
    styles = doc.styles
    
    # Title style
    try:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    except:
        title_style = styles['CustomTitle']
    
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(24)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(30)
    
    # Main heading style (16px)
    try:
        main_heading_style = styles.add_style('MainHeading', WD_STYLE_TYPE.PARAGRAPH)
    except:
        main_heading_style = styles['MainHeading']
    
    main_heading_font = main_heading_style.font
    main_heading_font.name = 'Times New Roman'
    main_heading_font.size = Pt(16)
    main_heading_font.bold = True
    main_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_heading_style.paragraph_format.space_before = Pt(40)
    main_heading_style.paragraph_format.space_after = Pt(20)
    main_heading_style.paragraph_format.page_break_before = True
    
    # Secondary heading style (14px)
    try:
        secondary_heading_style = styles.add_style('SecondaryHeading', WD_STYLE_TYPE.PARAGRAPH)
    except:
        secondary_heading_style = styles['SecondaryHeading']
    
    secondary_heading_font = secondary_heading_style.font
    secondary_heading_font.name = 'Times New Roman'
    secondary_heading_font.size = Pt(14)
    secondary_heading_font.bold = True
    secondary_heading_style.paragraph_format.space_before = Pt(25)
    secondary_heading_style.paragraph_format.space_after = Pt(15)
    
    # Tertiary heading style (12px)
    try:
        tertiary_heading_style = styles.add_style('TertiaryHeading', WD_STYLE_TYPE.PARAGRAPH)
    except:
        tertiary_heading_style = styles['TertiaryHeading']
    
    tertiary_heading_font = tertiary_heading_style.font
    tertiary_heading_font.name = 'Times New Roman'
    tertiary_heading_font.size = Pt(12)
    tertiary_heading_font.bold = True
    tertiary_heading_style.paragraph_format.space_before = Pt(20)
    tertiary_heading_style.paragraph_format.space_after = Pt(10)
    
    # Body text style (12px) - Reduced line height
    try:
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
    except:
        body_style = styles['BodyText']
    
    body_font = body_style.font
    body_font.name = 'Times New Roman'
    body_font.size = Pt(12)
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_style.paragraph_format.space_after = Pt(3)  # Further reduced
    body_style.paragraph_format.line_spacing = 1.0   # Single line spacing
    
    # TOC style
    try:
        toc_style = styles.add_style('TOCEntry', WD_STYLE_TYPE.PARAGRAPH)
    except:
        toc_style = styles['TOCEntry']
    
    toc_font = toc_style.font
    toc_font.name = 'Times New Roman'
    toc_font.size = Pt(12)
    toc_style.paragraph_format.space_after = Pt(3)  # Reduced spacing
    toc_style.paragraph_format.line_spacing = 1.0   # Single line spacing

def add_page_numbers(doc):
    """Add page numbers to footer - fixed implementation"""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Clear existing content
        footer_para.clear()
        
        # Add page number
        run = footer_para.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Add actual page number field code
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

def add_title_page(doc):
    """Add the title page"""
    title = doc.add_paragraph()
    title.style = 'CustomTitle'
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = "LUMIA ROBO-ADVISOR"
    
    subtitle1 = doc.add_paragraph("Intelligent Portfolio Management Platform", style='BodyText')
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle1.runs[0].font.size = Pt(18)
    subtitle1.runs[0].font.italic = True
    
    subtitle2 = doc.add_paragraph("Complete Project Documentation", style='BodyText')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(18)
    subtitle2.runs[0].font.italic = True
    
    # Add some space
    for _ in range(5):
        doc.add_paragraph()
    
    # Project details
    details = [
        "Project Type: Full-Stack Financial Technology Platform",
        "Technologies: React 18, TypeScript, Flask, PostgreSQL, Supabase",
        "Domain: Robo-Advisory and Portfolio Optimization",
        "Date: October 2025"
    ]
    
    for detail in details:
        p = doc.add_paragraph(detail, style='BodyText')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
    
    # Add page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents in proper tabular format with borders"""
    toc_title = doc.add_paragraph("TABLE OF CONTENTS", style='MainHeading')
    toc_title.paragraph_format.page_break_before = False
    
    # Create table for TOC with borders
    toc_table = doc.add_table(rows=1, cols=2)
    toc_table.style = 'Table Grid'
    
    # Header row
    header_cells = toc_table.rows[0].cells
    header_cells[0].text = "CONTENT"
    header_cells[1].text = "PAGE"
    
    # Format header
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    # Set column widths
    toc_table.columns[0].width = Inches(5.0)
    toc_table.columns[1].width = Inches(1.0)
    
    toc_entries = [
        ("Chapter 1: Introduction", "3"),
        ("1.1 Background", "3"),
        ("1.2 Objectives", "4"),
        ("1.3 Purpose, Scope, and Applicability", "5"),
        ("Chapter 2: Survey of Technologies", "8"),
        ("2.1 Technology Overview", "8"),
        ("2.1.1 Frontend Technologies", "9"),
        ("2.1.2 Backend Technologies", "12"),
        ("2.1.3 Database and Cloud Infrastructure", "15"),
        ("Chapter 3: Requirements and Analysis", "18"),
        ("3.1 Problem Definition", "18"),
        ("3.2 Requirements Specification", "21"),
        ("3.3 Analysis and Modeling", "25"),
        ("3.3.4 Gantt Chart", "27"),
        ("Chapter 4: System Design", "29"),
        ("4.1 Basic Modules", "29"),
        ("4.1 Use Case Diagram", "32"),
        ("4.2 Data Design", "33"),
        ("4.2.1 Schema Design", "33"),
        ("4.3 Data Integrity and Constraints - Class Diagram", "36"),
        ("4.4.1 UI Modules - Sequence Diagram", "38"),
        ("Chapter 5: Implementation and Testing", "40"),
        ("5.1 Implementation Approaches", "40"),
        ("5.2 Coding Details and Code Efficiency", "43"),
        ("Chapter 6: Results and Discussion", "47"),
        ("6.1 Test Reports", "47"),
        ("Chapter 7: Conclusions", "52"),
        ("7.1 Conclusion", "52"),
        ("7.2 Future Scope of the Project", "55")
    ]
    
    for entry, page in toc_entries:
        row = toc_table.add_row()
        
        # Content cell
        content_cell = row.cells[0]
        content_para = content_cell.paragraphs[0]
        content_para.paragraph_format.space_after = Pt(0)
        content_para.paragraph_format.line_spacing = 1.0
        
        if entry.startswith("Chapter"):
            run = content_para.add_run(entry)
            run.bold = True
        else:
            run = content_para.add_run("  " + entry)  # Indent sub-items
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Page number cell
        page_cell = row.cells[1]
        page_para = page_cell.paragraphs[0]
        page_para.paragraph_format.space_after = Pt(0)
        page_para.paragraph_format.line_spacing = 1.0
        page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_run = page_para.add_run(page)
        page_run.font.name = 'Times New Roman'
        page_run.font.size = Pt(12)
        
        # Set cell vertical alignment
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        page_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_page_break()

def add_actual_image(doc, image_path, title, description, width_inches=6):
    """Add actual image if available, otherwise placeholder"""
    # Add some space before
    doc.add_paragraph()
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    # Try to add actual image
    if os.path.exists(image_path):
        try:
            image_para = doc.add_paragraph()
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = image_para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            
            # Add description below image
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_para.add_run(description)
            desc_run.font.name = 'Times New Roman'
            desc_run.font.size = Pt(10)
            desc_run.font.italic = True
            
            print(f"✅ Added image: {image_path}")
            
        except Exception as e:
            print(f"❌ Could not add image {image_path}: {e}")
            # Fall back to placeholder
            add_image_placeholder_bordered(doc, title, description)
    else:
        print(f"📋 Image not found, using placeholder: {image_path}")
        # Fall back to placeholder
        add_image_placeholder_bordered(doc, title, description)
    
    # Add some space after
    doc.add_paragraph()

def add_image_placeholder_bordered(doc, title, description):
    """Add a bordered placeholder for missing images"""
    # Create a bordered table for image placeholder
    placeholder_table = doc.add_table(rows=1, cols=1)
    placeholder_table.style = 'Table Grid'
    
    cell = placeholder_table.cell(0, 0)
    cell.width = Inches(6)
    cell.height = Inches(3)
    
    # Add content to cell
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_para.paragraph_format.space_before = Pt(50)
    cell_para.paragraph_format.space_after = Pt(50)
    
    # Add placeholder text
    placeholder_run = cell_para.add_run(f"[{title}]")
    placeholder_run.font.name = 'Times New Roman'
    placeholder_run.font.size = Pt(12)
    placeholder_run.font.bold = True
    
    # Add line break and description
    cell_para.add_run("\n\n")
    desc_run = cell_para.add_run(description)
    desc_run.font.name = 'Times New Roman'
    desc_run.font.size = Pt(10)
    desc_run.font.italic = True
    
    # Center the cell content vertically
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_image_placeholder(doc, title, description):
    """Add image placeholder with border"""
    # Create a bordered table to simulate image placeholder
    placeholder_table = doc.add_table(rows=1, cols=1)
    placeholder_table.style = 'Table Grid'
    
    cell = placeholder_table.cell(0, 0)
    cell.width = Inches(6)
    
    # Add content to cell
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title
    title_run = cell_para.add_run(f"[{title}]")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # Add line break
    cell_para.add_run("\n")
    
    # Add description
    desc_run = cell_para.add_run(description)
    desc_run.font.name = 'Times New Roman'
    desc_run.font.size = Pt(10)
    desc_run.font.italic = True
    
    # Set minimum height
    cell._element.get_or_add_tcPr().append(
        OxmlElement('w:tcW')
    )

def process_text_content(text):
    """Process text content to identify headings and structure"""
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Remove ** formatting from text
        line = line.replace('**', '')
        
        # Identify different heading levels
        if line.startswith('CHAPTER ') and ':' in line:
            processed_lines.append(('main_heading', line))
        elif line.count('.') == 1 and line.split('.')[0].isdigit() and line.split('.')[1].strip().split()[0].isdigit():
            # Pattern like "1.1 Background"
            processed_lines.append(('secondary_heading', line))
        elif line.count('.') == 2 and all(part.strip().split()[0].isdigit() if part.strip() else False for part in line.split('.')[:2]):
            # Pattern like "1.1.1 Purpose"
            processed_lines.append(('tertiary_heading', line))
        elif line.endswith(':') and len(line.split()) <= 5:
            # Short lines ending with colon as subheadings
            processed_lines.append(('tertiary_heading', line))
        else:
            processed_lines.append(('body', line))
    
    return processed_lines

def read_chapter_file(filepath):
    """Read and return chapter content"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(filepath, 'r', encoding='latin-1') as f:
            return f.read()

def add_chapter_content(doc, chapter_num, chapter_file):
    """Add chapter content to document"""
    if not os.path.exists(chapter_file):
        print(f"Warning: Chapter file {chapter_file} not found")
        return
    
    content = read_chapter_file(chapter_file)
    processed_lines = process_text_content(content)
    
    for line_type, line_text in processed_lines:
        if line_type == 'main_heading':
            p = doc.add_paragraph(line_text, style='MainHeading')
        elif line_type == 'secondary_heading':
            p = doc.add_paragraph(line_text, style='SecondaryHeading')
            
            # Add Gantt Chart after 3.3.4 heading
            if "3.3.4" in line_text and "Gantt" in line_text:
                add_gantt_chart_image(doc)
            
            # Add Use Case Diagram after 4.1 Basic Modules heading
            elif "4.1" in line_text and "Basic Modules" in line_text:
                add_use_case_diagram_image(doc)
            
            # Add Schema Design after 4.2.1 Schema Design heading
            elif "4.2.1" in line_text and "Schema Design" in line_text:
                add_schema_design_image(doc)
                
        elif line_type == 'tertiary_heading':
            p = doc.add_paragraph(line_text, style='TertiaryHeading')
            
            # Add Class Diagram after "Data Integrity and Constraints" heading
            if "Data Integrity and Constraints" in line_text:
                add_class_diagram_image(doc)
            
            # Add Sequence Diagram after "4.4.1 UI Modules" heading
            elif "4.4.1" in line_text and "UI Modules" in line_text:
                add_sequence_diagram_image(doc)
                
        else:
            if line_text.strip():
                p = doc.add_paragraph(line_text, style='BodyText')

def add_gantt_chart_image(doc):
    """Add Gantt Chart image at 3.3.4"""
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    gantt_paths = [
        os.path.join(blackbook_dir, "images", "gantt_chart.png"),
        os.path.join(blackbook_dir, "images", "gantt_chart.jpg"),
        os.path.join(blackbook_dir, "images", "gantt.png"),
        os.path.join(blackbook_dir, "images", "timeline.png"),
        os.path.join(blackbook_dir, "images", "project_timeline.png"),
    ]
    
    image_found = False
    for path in gantt_paths:
        if os.path.exists(path):
            add_actual_image(doc, path, "GANTT CHART - PROJECT TIMELINE", "Comprehensive Development Schedule", 7)
            image_found = True
            break
    
    if not image_found:
        add_image_placeholder_bordered(doc, "GANTT CHART - PROJECT TIMELINE", "Comprehensive Development Schedule")

def add_use_case_diagram_image(doc):
    """Add Use Case Diagram image at 4.1 Basic Modules"""
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    use_case_paths = [
        os.path.join(blackbook_dir, "images", "use_case.png"),
        os.path.join(blackbook_dir, "images", "use_case_diagram.png"),
        os.path.join(blackbook_dir, "images", "usecase.png"),
        os.path.join(blackbook_dir, "images", "use_cases.png"),
    ]
    
    image_found = False
    for path in use_case_paths:
        if os.path.exists(path):
            add_actual_image(doc, path, "USE CASE DIAGRAM", "Lumia Robo-Advisor System Use Cases", 7)
            image_found = True
            break
    
    if not image_found:
        add_image_placeholder_bordered(doc, "USE CASE DIAGRAM", "Lumia Robo-Advisor System Use Cases")

def add_schema_design_image(doc):
    """Add Schema Design image at 4.2.1 Schema Design"""
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    schema_paths = [
        os.path.join(blackbook_dir, "images", "schema.png"),
        os.path.join(blackbook_dir, "images", "database_schema.png"),
        os.path.join(blackbook_dir, "images", "db_schema.png"),
        os.path.join(blackbook_dir, "images", "erd.png"),
        os.path.join(blackbook_dir, "images", "database.png"),
    ]
    
    image_found = False
    for path in schema_paths:
        if os.path.exists(path):
            add_actual_image(doc, path, "DATABASE SCHEMA DESIGN", "Entity Relationship Diagram", 8)
            image_found = True
            break
    
    if not image_found:
        add_image_placeholder_bordered(doc, "DATABASE SCHEMA DESIGN", "Entity Relationship Diagram")

def add_class_diagram_image(doc):
    """Add Class Diagram image after Data Integrity and Constraints"""
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    class_paths = [
        os.path.join(blackbook_dir, "images", "class_diagram.png"),
        os.path.join(blackbook_dir, "images", "class.png"),
        os.path.join(blackbook_dir, "images", "component_architecture.png"),
        os.path.join(blackbook_dir, "images", "architecture.png"),
    ]
    
    image_found = False
    for path in class_paths:
        if os.path.exists(path):
            add_actual_image(doc, path, "CLASS DIAGRAM", "Data Integrity and Constraints Design", 8)
            image_found = True
            break
    
    if not image_found:
        add_image_placeholder_bordered(doc, "CLASS DIAGRAM", "Data Integrity and Constraints Design")

def add_sequence_diagram_image(doc):
    """Add Sequence Diagram image at 4.4.1 UI Modules"""
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    sequence_paths = [
        os.path.join(blackbook_dir, "images", "sequence_diagram.png"),
        os.path.join(blackbook_dir, "images", "sequence.png"),
        os.path.join(blackbook_dir, "images", "portfolio_sequence.png"),
        os.path.join(blackbook_dir, "images", "ui_sequence.png"),
    ]
    
    image_found = False
    for path in sequence_paths:
        if os.path.exists(path):
            add_actual_image(doc, path, "SEQUENCE DIAGRAM", "UI Modules Interaction Sequence", 8)
            image_found = True
            break
    
    if not image_found:
        add_image_placeholder_bordered(doc, "SEQUENCE DIAGRAM", "UI Modules Interaction Sequence")

def main():
    """Main function to create the DOCX document"""
    print("Creating Lumia Black Book DOCX document...")
    
    # Create new document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Add title page
    add_title_page(doc)
    
    # Add table of contents
    add_table_of_contents(doc)
    
    # Add page numbers (starting from content pages)
    add_page_numbers(doc)
    
    # Process chapters 1-7 ONLY (NO CHAPTER 0)
    chapter_files = [
        ('chapter1.txt', 1),
        ('chapter2.txt', 2),
        ('chapter3.txt', 3),
        ('chapter4.txt', 4),
        ('chapter5.txt', 5),
        ('chapter6.txt', 6),
        ('chapter7.txt', 7)
    ]
    
    blackbook_dir = os.path.dirname(os.path.abspath(__file__))
    creation_files_dir = os.path.join(blackbook_dir, 'blackbookcreationfiles')
    
    for chapter_file, chapter_num in chapter_files:
        chapter_path = os.path.join(creation_files_dir, chapter_file)
        print(f"Processing {chapter_file}...")
        add_chapter_content(doc, chapter_num, chapter_path)
    
    # Save document
    output_path = os.path.join(blackbook_dir, 'Lumia_BlackBook_With_Images.docx')
    doc.save(output_path)
    
    print(f"Black Book DOCX created successfully: {output_path}")
    print("\nDocument specifications:")
    print("- Font: Times New Roman")
    print("- Main headings: 16px")
    print("- Secondary headings: 14px") 
    print("- Body text: 12px")
    print("- Single line spacing (1.0) for compact formatting")
    print("- Minimal paragraph spacing (3pt)")
    print("- Table of Contents in clean tabular format")
    print("- Page numbers added to all content pages")
    print("- ** formatting removed from text")
    print("- Actual images included where available")
    print("- Professional compact formatting")

if __name__ == "__main__":
    main()