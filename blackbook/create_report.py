import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

def get_style_from_html(html_string):
    """Parses an HTML string to extract CSS styles."""
    style = {}
    if not html_string:
        return style

    soup = BeautifulSoup(html_string, 'html.parser')
    div = soup.find('div')
    if not div or not div.get('style'):
        return style

    style_str = div.get('style')
    
    # Extract font-size
    size_match = re.search(r'font-size:\s*(\d+)px', style_str)
    if size_match:
        style['font_size'] = int(size_match.group(1))

    # Extract font-weight
    if 'font-weight: bold' in style_str:
        style['bold'] = True

    # Extract text-align
    align_match = re.search(r'text-align:\s*(\w+)', style_str)
    if align_match:
        style['align'] = align_match.group(1)
        
    # Extract text-decoration
    if 'text-decoration: underline' in style_str:
        style['underline'] = True
        
    # Extract margin-left for indentation
    margin_match = re.search(r'margin-left:\s*(\d+)px', style_str)
    if margin_match:
        style['indent'] = int(margin_match.group(1))

    return style

def parse_bold_text(text):
    """Parse text with **bold** markers and return segments."""
    segments = []
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Add text before bold
        if match.start() > last_end:
            segments.append({'text': text[last_end:match.start()], 'bold': False})
        # Add bold text
        segments.append({'text': match.group(1), 'bold': True})
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        segments.append({'text': text[last_end:], 'bold': False})
    
    return segments if segments else [{'text': text, 'bold': False}]

def add_formatted_paragraph(document, text, style_dict=None, is_code=False):
    """Add a paragraph with proper formatting including bold text."""
    p = document.add_paragraph()
    
    if style_dict and style_dict.get('indent'):
        p.paragraph_format.left_indent = Inches(style_dict.get('indent') / 72.0)
    
    segments = parse_bold_text(text)
    
    for segment in segments:
        run = p.add_run(segment['text'])
        
        if is_code:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            run._element.get_or_add_rPr().append(shading_elm)
        else:
            run.font.name = 'Times New Roman'
            if style_dict and style_dict.get('font_size'):
                run.font.size = Pt(style_dict.get('font_size'))
            else:
                run.font.size = Pt(12)
        
        if segment['bold'] or (style_dict and style_dict.get('bold')):
            run.bold = True
            
        if style_dict and style_dict.get('underline'):
            run.underline = True
    
    # Set alignment
    if style_dict and style_dict.get('align'):
        alignment = style_dict.get('align')
        if alignment == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return p

def parse_table(lines, start_idx):
    """Parse markdown table and return rows and number of lines consumed."""
    rows = []
    i = start_idx
    
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        # Skip separator lines like |:---:|:---|:---:|
        if re.match(r'^\|[\s\:\-\|]+\|$', line):
            i += 1
            continue
        
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    
    return rows, i - start_idx

def add_table_to_doc(document, rows):
    """Add a properly formatted table to the document."""
    if not rows:
        return
    
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # Parse HTML if present
            soup = BeautifulSoup(cell_text, 'html.parser')
            text = soup.get_text().strip()
            
            # Add formatted text
            paragraph = cell.paragraphs[0]
            paragraph.text = ''
            segments = parse_bold_text(text)
            
            for segment in segments:
                run = paragraph.add_run(segment['text'])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if segment['bold'] or i == 0:  # Make header row bold
                    run.bold = True

def convert_md_to_docx(source_dir, output_file):
    """Converts markdown files with embedded HTML styling to a DOCX document."""
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Modify heading styles to be black (not blue)
    for i in range(1, 10):
        try:
            heading_style = document.styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
        except:
            pass

    # Get chapter files in order
    files = sorted([f for f in os.listdir(source_dir) if f.startswith('chapter') and f.endswith('.md')])

    for filename in files:
        filepath = os.path.join(source_dir, filename)
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Handle code blocks
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    # End of code block - add all lines
                    for code_line in code_block_lines:
                        add_formatted_paragraph(document, code_line, is_code=True)
                    in_code_block = False
                    code_block_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue
            
            if not line:
                i += 1
                continue

            # Handle HTML div with styling
            if line.startswith('<div'):
                soup = BeautifulSoup(line, 'html.parser')
                text = soup.get_text().strip()
                style_dict = get_style_from_html(line)
                
                if text.startswith('#'):
                    # It's a heading
                    level = text.count('#')
                    text = text.lstrip('# ').strip()
                    p = document.add_heading(text, level=level)
                    if style_dict.get('align') == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if style_dict.get('underline'):
                        for run in p.runs:
                            run.underline = True
                elif text and text != '':
                    add_formatted_paragraph(document, text, style_dict)
                
                i += 1
                continue

            # Handle tables
            if line.startswith('|'):
                rows, lines_consumed = parse_table(lines, i)
                add_table_to_doc(document, rows)
                i += lines_consumed
                continue
            
            # Handle horizontal rules
            if line == '---':
                document.add_paragraph('_' * 50)
                i += 1
                continue
            
            # Handle markdown headings
            if line.startswith('#'):
                level = line.count('#', 0, 6)
                text = line.lstrip('# ').strip()
                document.add_heading(text, level=level)
                i += 1
                continue
            
            # Handle bullet points
            if line.startswith('-') or line.startswith('*'):
                text = line.lstrip('-* ').strip()
                add_formatted_paragraph(document, text, {'indent': 20})
                i += 1
                continue
            
            # Handle numbered lists
            if re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                add_formatted_paragraph(document, text, {'indent': 20})
                i += 1
                continue
            
            # Regular paragraph
            if line and not line.startswith('<'):
                add_formatted_paragraph(document, line)
            
            i += 1
        
        document.add_page_break()

    print(f"Saving document to {output_file}")
    document.save(output_file)
    print("Done!")

if __name__ == '__main__':
    # Ensure you have the required packages:
    # pip install python-docx beautifulsoup4
    
    blackbook_dir = os.path.dirname(__file__)
    output_filename = os.path.join(os.path.dirname(__file__), '..', 'Lumia_Project_Report.docx')
    
    convert_md_to_docx(blackbook_dir, output_filename)
