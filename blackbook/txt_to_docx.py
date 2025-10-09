"""
Convert plain text files to a Word document
Creates a professional Word document from clean text files
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def format_title_page(doc, lines, start_idx):
    """Format the title page section"""
    idx = start_idx
    
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line:
            idx += 1
            continue
            
        # Stop at certificate section
        if 'BHAVNA TRUST JUNIOR AND DEGREE COLLEGE' in line and idx > 10:
            break
            
        # Main title (16px)
        if 'PROJECT REPORT' in line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
        
        # Project title (16px)
        elif 'Lumia Robo-Advisor' in line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.bold = True
        
        # Other centered text
        elif any(x in line for x in ['SUBMITTED TO', 'DEPARTMENT', 'COLLEGE', 'MUMBAI', 
                                      'University', 'Fulfilment', 'Bachelor', 'COMPUTER SCIENCE',
                                      'SUBMITTED BY', 'MAHALDAR']):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if any(x in line for x in ['MAHALDAR', 'DEPARTMENT']):
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(14)
            else:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
        
        # Guide names
        elif 'HEAD OF DEPARTMENT' in line or 'PROJECT GUIDE' in line:
            # Split the line if it contains both
            if 'HEAD OF DEPARTMENT' in line and 'PROJECT GUIDE' in line:
                parts = line.split('PROJECT GUIDE')
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = p.add_run(parts[0].replace('HEAD OF DEPARTMENT', 'HEAD OF DEPARTMENT: '))
                run2 = p.add_run('          ')
                run3 = p.add_run('PROJECT GUIDE: ' + parts[1])
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
        
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
        
        idx += 1
    
    return idx

def format_section(doc, title, content_lines):
    """Format a section with title and content"""
    # Add section title (Main heading - 16px)
    p = doc.add_heading(title, level=1)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Check if this is the INDEX section
    if title == 'INDEX':
        # Parse and add table
        table_rows = parse_table_from_text(content_lines)
        
        if table_rows:
            # Add the table
            add_table_to_doc(doc, table_rows)
        
        # Also add non-table content
        for line in content_lines:
            line_stripped = line.strip()
            
            # Skip table rows
            if line_stripped.startswith('|') or '---' in line_stripped:
                continue
            
            # Add chapter headings
            if 'CHAPTER' in line_stripped and ':' in line_stripped:
                p = doc.add_paragraph(line_stripped)
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line_stripped and line_stripped != 'End of Index':
                p = doc.add_paragraph(line_stripped)
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        # Add content for other sections
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            
            # Regular paragraph (12px)
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

def parse_table_from_text(lines):
    """Parse table from text lines"""
    rows = []
    in_table = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Detect table rows
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            # Skip separator lines
            if '---' in line_stripped or ':---:' in line_stripped:
                in_table = True
                continue
            
            # Parse the row
            cells = [cell.strip() for cell in line_stripped.split('|')[1:-1]]
            if cells and any(cells):  # Only add non-empty rows
                rows.append(cells)
                in_table = True
        elif in_table and line_stripped:
            # Check if it's a plain table row without pipes
            if any(keyword in line_stripped for keyword in ['CHAPTER', '|']):
                continue
    
    return rows

def add_table_to_doc(doc, rows):
    """Add a table to the document"""
    if not rows or len(rows) < 2:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    # Fill the table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            
            # Format all cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    
                    # Format header row
                    if i == 0:
                        run.font.bold = True
            
            # Add borders to all cells
            set_cell_border(cell, top={}, bottom={}, left={}, right={})

def convert_txt_to_docx(txt_dir, output_file):
    """Convert text files to a single Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Get all text files in order
    txt_files = [
        'chapter0.txt',
        'chapter1.txt',
        'chapter2.txt',
        'chapter3.txt',
        'chapter4.txt',
        'chapter5.txt',
        'chapter6.txt',
        'chapter7.txt'
    ]
    
    print("Converting text files to Word document...")
    
    for txt_file in txt_files:
        txt_path = os.path.join(txt_dir, txt_file)
        
        if not os.path.exists(txt_path):
            print(f"  ⚠ Skipping {txt_file} (not found)")
            continue
        
        print(f"  Processing {txt_file}...")
        
        # Read the text file
        with open(txt_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Process based on chapter
        if txt_file == 'chapter0.txt':
            # Special handling for chapter 0 (front matter)
            idx = 0
            current_section = []
            current_title = None
            
            for line in lines:
                line_stripped = line.strip()
                
                # Detect section titles
                if line_stripped in ['DECLARATION', 'ACKNOWLEDGEMENT', 'ABSTRACT', 'INDEX']:
                    # Add previous section
                    if current_title and current_section:
                        format_section(doc, current_title, current_section)
                        add_page_break(doc)
                    
                    current_title = line_stripped
                    current_section = []
                
                elif line_stripped == 'Certificate':
                    if current_section:
                        add_page_break(doc)
                    current_title = 'CERTIFICATE'
                    current_section = []
                
                elif not current_title and idx < 30:
                    # Title page content
                    continue
                
                else:
                    current_section.append(line)
            
            # Add last section
            if current_title and current_section:
                format_section(doc, current_title, current_section)
        
        else:
            # Regular chapters
            chapter_num = txt_file.replace('chapter', '').replace('.txt', '')
            
            # Add chapter heading (16px)
            if lines:
                first_line = lines[0].strip()
                if first_line:
                    heading = doc.add_heading(first_line, level=1)
                    heading.runs[0].font.name = 'Times New Roman'
                    heading.runs[0].font.size = Pt(16)
                    heading.runs[0].font.bold = True
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                
                # Add rest of content
                for line in lines[1:]:
                    line_stripped = line.strip()
                    
                    if not line_stripped:
                        continue
                    
                    # Check for subheadings (14px)
                    if any(x in line_stripped for x in ['Background', 'Objectives', 'Purpose', 
                                                         'Problem Definition', 'Requirements',
                                                         'Basic Modules', 'Data Design',
                                                         'Implementation', 'Testing', 'Conclusion',
                                                         'Scope', 'Applicability']):
                        h = doc.add_heading(line_stripped, level=2)
                        h.runs[0].font.name = 'Times New Roman'
                        h.runs[0].font.size = Pt(14)
                        h.runs[0].font.bold = True
                        h.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        # Regular paragraph (12px)
                        p = doc.add_paragraph(line_stripped)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
        
        # Add page break after each chapter (except the last)
        if txt_file != txt_files[-1]:
            add_page_break(doc)
    
    # Save document
    print(f"\nSaving document to {output_file}")
    doc.save(output_file)
    print("Done!")

if __name__ == "__main__":
    # Get the directory where this script is located
    script_dir = os.path.dirname(os.path.abspath(__file__))
    txt_dir = os.path.join(script_dir, 'txt_files')
    output_file = os.path.join(script_dir, '..', 'Lumia_Project_Report_Clean.docx')
    
    print("=" * 60)
    print("Converting Text Files to Word Document")
    print("=" * 60)
    
    convert_txt_to_docx(txt_dir, output_file)
    
    print("\n" + "=" * 60)
    print("Conversion Complete!")
    print(f"Output: {os.path.abspath(output_file)}")
    print("=" * 60)
