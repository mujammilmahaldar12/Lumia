"""
Lumia DOCX Generator
Converts Lumia text files into a professionally formatted DOCX document
Created by: Mujammil Mahaldar
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

class LumiaDocxGenerator:
    def __init__(self, output_filename="Lumia_Financial_Analytics_Platform.docx"):
        self.doc = Document()
        self.output_filename = output_filename
        self.setup_styles()
        
    def setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Title Style (18pt)
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Main Heading Style (16pt)
        if 'Main Heading' not in [s.name for s in styles]:
            main_heading_style = styles.add_style('Main Heading', WD_STYLE_TYPE.PARAGRAPH)
            main_font = main_heading_style.font
            main_font.name = 'Calibri'
            main_font.size = Pt(16)
            main_font.bold = True
            main_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            main_heading_style.paragraph_format.space_before = Pt(12)
            main_heading_style.paragraph_format.space_after = Pt(6)
        
        # Sub Heading Style (14pt)
        if 'Sub Heading' not in [s.name for s in styles]:
            sub_heading_style = styles.add_style('Sub Heading', WD_STYLE_TYPE.PARAGRAPH)
            sub_font = sub_heading_style.font
            sub_font.name = 'Calibri'
            sub_font.size = Pt(14)
            sub_font.bold = True
            sub_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_heading_style.paragraph_format.space_before = Pt(8)
            sub_heading_style.paragraph_format.space_after = Pt(4)
        
        # Body Text Style (12pt)
        if 'Body Text' not in [s.name for s in styles]:
            body_style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(14)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.15
        
        # Code Style (12pt)
        if 'Code Style' not in [s.name for s in styles]:
            code_style = styles.add_style('Code Style', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(12)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_after = Pt(6)
    
    def add_page_break(self):
        """Add a page break"""
        self.doc.add_page_break()
    
    def process_text_formatting(self, text):
        """Process text to identify headings, subheadings, and body text"""
        lines = text.split('\n')
        processed_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                processed_content.append(('empty', ''))
                continue
            
            # Check for different heading levels
            if line.startswith('CHAPTER') or (line.upper() == line and len(line) > 10 and '**' not in line):
                processed_content.append(('title', line))
            elif re.match(r'^\d+\.\d+\.\d+', line):
                processed_content.append(('sub_heading', line))
            elif re.match(r'^\d+\.\d+', line) and '**' not in line:
                processed_content.append(('sub_heading', line))
            elif re.match(r'^\d+\.', line) and '**' not in line:
                processed_content.append(('main_heading', line))
            elif line.endswith(':') and '**' not in line:
                processed_content.append(('sub_heading', line))
            elif line.startswith('```') or line.startswith('class ') or line.startswith('def ') or line.startswith('import '):
                processed_content.append(('code', line))
            elif line.startswith('-') or line.startswith('•'):
                processed_content.append(('bullet', line))
            else:
                # Always check for ** formatting for any remaining text
                processed_content.append(('mixed_bold', line))
        
        return processed_content
    
    def add_formatted_content(self, content_type, text, is_chapter_title=False):
        """Add content with appropriate formatting"""
        if content_type == 'empty':
            self.doc.add_paragraph()
        elif content_type == 'title':
            p = self.doc.add_paragraph(text, style='Custom Title')
            if is_chapter_title:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif content_type == 'main_heading':
            # Handle headings that might have ** in them
            p = self.doc.add_paragraph(style='Main Heading')
            self.add_mixed_bold_text(p, text)
        elif content_type == 'sub_heading':
            # Handle sub-headings that might have ** in them
            p = self.doc.add_paragraph(style='Sub Heading')
            self.add_mixed_bold_text(p, text)
        elif content_type == 'bold':
            # Create paragraph with entirely bold text
            p = self.doc.add_paragraph(style='Body Text')
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        elif content_type == 'mixed_bold':
            # Handle inline bold formatting in text
            p = self.doc.add_paragraph(style='Body Text')
            self.add_mixed_bold_text(p, text)
        elif content_type == 'code':
            p = self.doc.add_paragraph(text, style='Code Style')
        elif content_type == 'bullet':
            # Handle bullets with potential ** formatting
            bullet_text = text[1:].strip()
            p = self.doc.add_paragraph(style='Body Text')
            p.paragraph_format.left_indent = Inches(0.5)
            self.add_mixed_bold_text(p, bullet_text)
        else:  # body text
            p = self.doc.add_paragraph(style='Body Text')
            self.add_mixed_bold_text(p, text)
    
    def add_mixed_bold_text(self, paragraph, text):
        """Add text with mixed bold and normal formatting - COMPLETELY REWRITTEN"""
        # Simple approach: replace all ** patterns first, then add formatting
        import re
        
        # Pattern to find **text** 
        pattern = r'\*\*(.*?)\*\*'
        
        last_end = 0
        
        # Find all matches of **text**
        for match in re.finditer(pattern, text):
            # Add normal text before the match
            if match.start() > last_end:
                normal_text = text[last_end:match.start()]
                if normal_text:
                    run = paragraph.add_run(normal_text)
                    run.font.size = Pt(12)
            
            # Add bold text (without the **)
            bold_text = match.group(1)
            if bold_text:
                run = paragraph.add_run(bold_text)
                run.bold = True
                run.font.size = Pt(12)
            
            last_end = match.end()
        
        # Add remaining text after last match
        if last_end < len(text):
            remaining_text = text[last_end:]
            if remaining_text:
                run = paragraph.add_run(remaining_text)
                run.font.size = Pt(12)
    
    def add_cover_page(self, main_content):
        """Add a professional cover page"""
        # Title
        title = self.doc.add_paragraph('LUMIA AI FINANCIAL ANALYTICS PLATFORM', style='Custom Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.doc.add_paragraph('A Project Report', style='Sub Heading')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Academic info
        academic_info = [
            "Submitted in partial fulfillment of the",
            "requirements for the award of the Degree of",
            "BACHELOR OF SCIENCE (INFORMATION TECHNOLOGY)",
            "(Semester V Project)"
        ]
        
        for info in academic_info:
            p = self.doc.add_paragraph(info, style='Body Text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Student info
        student_p = self.doc.add_paragraph('By', style='Body Text')
        student_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        name_p = self.doc.add_paragraph('MUJAMMIL MAHALDAR', style='Sub Heading')
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # College info
        college_info = [
            "DEPARTMENT OF INFORMATION TECHNOLOGY",
            "BHAVNA TRUST'S JUNIOR & DEGREE COLLEGE OF COMMERCE & SCIENCE",
            "(Affiliated to University of Mumbai)",
            "MUMBAI - 400088 (MAHARASHTRA)",
            "Academic Year: 2024-2025"
        ]
        
        for info in college_info:
            p = self.doc.add_paragraph(info, style='Body Text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.add_page_break()
    
    def add_table_of_contents(self):
        """Add a table of contents"""
        toc_title = self.doc.add_paragraph('TABLE OF CONTENTS', style='Custom Title')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # TOC items
        toc_items = [
            ("Certificate", "i"),
            ("Acknowledgement", "ii"),
            ("Abstract", "iii"),
            ("Chapter 1: Introduction", "1"),
            ("Chapter 2: Literature Survey", "15"),
            ("Chapter 3: System Analysis and Design", "25"),
            ("Chapter 4: System Requirements", "35"),
            ("Chapter 5: Implementation and Testing", "45"),
            ("Chapter 6: Results and Conclusion", "65"),
            ("References", "75"),
            ("Appendix", "80")
        ]
        
        for item, page in toc_items:
            p = self.doc.add_paragraph()
            p.style = 'Body Text'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Keep TOC items left-aligned
            run1 = p.add_run(item)
            run1.font.size = Pt(12)  # Ensure 12pt font
            run2 = p.add_run(f" {'.' * (50 - len(item))} {page}")
            run2.font.name = 'Calibri'
            run2.font.size = Pt(12)  # Ensure 12pt font
        
        self.add_page_break()
    
    def read_file_content(self, filepath):
        """Read content from a file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()
        except FileNotFoundError:
            print(f"Warning: File {filepath} not found")
            return ""
        except Exception as e:
            print(f"Error reading {filepath}: {e}")
            return ""
    
    def generate_document(self):
        """Generate the complete DOCX document"""
        print("🔮 Generating Lumia Financial Analytics Platform Documentation...")
        
        # Base directory
        base_dir = r"c:\Users\Mujammil\Desktop\Lumia\app"
        
        # Read main content
        main_content = self.read_file_content(os.path.join(base_dir, "lumia.txt"))
        
        # Add cover page
        self.add_cover_page(main_content)
        
        # Add table of contents
        self.add_table_of_contents()
        
        # Process main content sections
        if main_content:
            sections = main_content.split("------------------------------------------------------------")
            for i, section in enumerate(sections):
                if section.strip():
                    if i > 0:  # Skip title section as it's already in cover page
                        processed_content = self.process_text_formatting(section.strip())
                        for content_type, text in processed_content:
                            self.add_formatted_content(content_type, text, is_chapter_title=False)
                        if i < len(sections) - 1:
                            self.doc.add_paragraph()
        
        # Add chapters
        for chapter_num in range(1, 7):
            chapter_file = os.path.join(base_dir, f"lumia_chapter{chapter_num}.txt")
            chapter_content = self.read_file_content(chapter_file)
            
            if chapter_content:
                # Start each chapter on a new page
                self.add_page_break()
                
                # Add chapter title with center alignment
                chapter_title = f"CHAPTER {chapter_num}"
                if chapter_num == 1:
                    chapter_title += " - INTRODUCTION"
                elif chapter_num == 2:
                    chapter_title += " - LITERATURE SURVEY"
                elif chapter_num == 3:
                    chapter_title += " - SYSTEM ANALYSIS AND DESIGN"
                elif chapter_num == 4:
                    chapter_title += " - SYSTEM REQUIREMENTS"
                elif chapter_num == 5:
                    chapter_title += " - IMPLEMENTATION AND TESTING"
                elif chapter_num == 6:
                    chapter_title += " - RESULTS AND CONCLUSION"
                
                # Add chapter title with center alignment
                chapter_p = self.doc.add_paragraph(chapter_title, style='Custom Title')
                chapter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.doc.add_paragraph()
                
                # Process chapter content
                processed_content = self.process_text_formatting(chapter_content)
                for content_type, text in processed_content:
                    self.add_formatted_content(content_type, text, is_chapter_title=False)
        
        print("✅ Document structure created successfully!")
    
    def save_document(self):
        """Save the document"""
        try:
            output_path = os.path.join(r"c:\Users\Mujammil\Desktop\Lumia", self.output_filename)
            self.doc.save(output_path)
            print(f"📄 Document saved successfully as: {output_path}")
            return output_path
        except Exception as e:
            print(f"❌ Error saving document: {e}")
            return None

def main():
    """Main function to generate the DOCX"""
    print("=" * 60)
    print("🔮 LUMIA FINANCIAL ANALYTICS PLATFORM")
    print("📚 Professional DOCX Generator")
    print("=" * 60)
    
    # Create generator instance with final bold test
    generator = LumiaDocxGenerator("Lumia_Final_Bold_Test.docx")
    
    # Generate document
    generator.generate_document()
    
    # Save document
    output_path = generator.save_document()
    
    if output_path:
        print("\n" + "=" * 60)
        print("🎉 SUCCESS! Your professional Lumia documentation is ready!")
        print(f"📁 Location: {output_path}")
        print("=" * 60)
        print("\n📋 Document Features:")
        print("✓ Professional cover page")
        print("✓ Table of contents")
        print("✓ Proper heading hierarchy (16pt → 14pt → 12pt)")
        print("✓ Formatted code sections")
        print("✓ Justified body text")
        print("✓ Page breaks between chapters")
        print("✓ Academic formatting standards")
        print("\n🚀 Ready for submission!")
    else:
        print("❌ Failed to generate document. Please check the error messages above.")

if __name__ == "__main__":
    main()